"""v4.1: integración con python-docx (preserva estilos del documento).

Estrategia:
1. Abre v3.18-aceptada.docx con python-docx
2. Para cada capítulo a reemplazar (Mensajes clave, RE, Caps 2-8):
   a. Encuentra el bloque (start_para, end_para)
   b. Borra todos los párrafos del bloque
   c. Inserta nuevos párrafos del .md correspondiente, usando los estilos
      Heading1/2/3 y Normal que ya están definidos en v3.18
3. Markdown inline: **bold** → bold real (sin asteriscos), *italic*, listas, tablas
4. Guarda como informe-final-v4.1.docx

Salida: informe-final-v4.1.docx (LIMPIA — sin tracked changes)
"""

from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from copy import deepcopy
from lxml import etree

BASE = Path("/srv/projects/cochid/cochid-scribe/docs/cif-review")
SRC = BASE / "output/informe-final-v3.18-aceptada.docx"
V41 = BASE / "output/informe-final-v4.1.docx"
CONTEXT = Path("/tmp/v40-context/output")


# ============================================================
# Inline markdown parsing (returns list of (text, bold, italic) tuples)
# ============================================================
def parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    """Parse **bold**, *italic*, `code` into runs."""
    # Strip leading/trailing whitespace mark
    runs = []
    pos = 0
    pattern = re.compile(r"\*\*([^*]+)\*\*|\*([^*\n]+)\*|`([^`]+)`")
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, False))
        if m.group(1):  # bold
            runs.append((m.group(1), True, False))
        elif m.group(2):  # italic
            runs.append((m.group(2), False, True))
        elif m.group(3):  # code
            runs.append((m.group(3), False, False))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    return runs if runs else [(text, False, False)]


# ============================================================
# Parse markdown file to blocks
# ============================================================
def parse_md(md_path: Path) -> list[dict]:
    if not md_path.exists():
        return []
    blocks = []
    text = md_path.read_text()
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if i < 5 and stripped.startswith("# "):
            i += 1
            continue
        m = re.match(r"^\[Heading(\d)\]\s*(.+)$", stripped)
        if m:
            blocks.append({"kind": "heading", "level": int(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue
        m = re.match(r"^\[p\]\s*(.+)$", stripped)
        if m:
            buf = [m.group(1)]
            j = i + 1
            while j < len(lines):
                nl = lines[j].rstrip()
                ns = nl.strip()
                if not ns:
                    j += 1
                    continue
                if re.match(r"^\[Heading\d\]|^\[p\]|^\|", ns):
                    break
                buf.append(ns)
                j += 1
            blocks.append({"kind": "p", "text": " ".join(buf)})
            i = j
            continue
        if stripped.startswith("|"):
            tbl_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                tbl_lines.append(lines[j].strip())
                j += 1
            rows = []
            for tl in tbl_lines:
                if re.match(r"^\|[\s:|-]+\|$", tl):
                    continue
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if rows:
                blocks.append({"kind": "table", "rows": rows})
            i = j
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            items = [stripped[2:].strip()]
            j = i + 1
            while j < len(lines):
                ns = lines[j].strip()
                if ns.startswith("- ") or ns.startswith("* "):
                    items.append(ns[2:].strip())
                    j += 1
                elif not ns:
                    j += 1
                else:
                    break
            blocks.append({"kind": "list", "items": items})
            i = j
            continue
        blocks.append({"kind": "p", "text": stripped})
        i += 1
    return blocks


def split_combined(blocks, sections):
    result = {label: [] for label, _ in sections}
    current_label = None
    for b in blocks:
        if b["kind"] == "heading" and b["level"] == 1:
            for label, matcher in sections:
                if matcher(b["text"]):
                    current_label = label
                    break
        if current_label:
            result[current_label].append(b)
    return result


# ============================================================
# Insert helpers using python-docx
# ============================================================
def add_paragraph_with_runs(doc, text, style=None):
    """Add paragraph with parsed inline markdown."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    runs = parse_inline(text)
    for run_text, bold, italic in runs:
        if not run_text:
            continue
        r = p.add_run(run_text)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
    return p


def add_table_to_doc(doc, rows, style="Table Grid"):
    """Add a table to the document."""
    if not rows:
        return None
    n_rows = len(rows)
    n_cols = len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        table.style = doc.styles[style]
    except KeyError:
        # Fallback: set borders manually via XML
        pass
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = ""  # clear
                # Use first run + parse inline
                p = cell.paragraphs[0]
                p.text = ""
                for run_text, bold, italic in parse_inline(cell_text):
                    if not run_text:
                        continue
                    r = p.add_run(run_text)
                    if bold:
                        r.bold = True
                    if italic:
                        r.italic = True
    return table


def insert_blocks_at(doc, anchor_para, blocks):
    """Insert blocks as new paragraphs/tables before anchor_para."""
    # We use the underlying XML element to insert
    body = doc.element.body
    anchor_el = anchor_para._element

    elements_added = 0
    # Get a doc to add to (we'll construct elements then move them)
    # Easier: append everything at end with python-docx, then move

    # Approach: create a temporary doc context for adding
    # We'll directly construct XML elements following v3.18 patterns

    for b in blocks:
        if b["kind"] == "heading":
            level = b["level"]
            style = f"Heading {level}"
            # Create heading paragraph
            p = doc.add_paragraph()
            try:
                p.style = doc.styles[style]
            except KeyError:
                # Fallback styleId
                p.style = doc.styles[f"Heading{level}"]
            for run_text, bold, italic in parse_inline(b["text"]):
                if not run_text:
                    continue
                r = p.add_run(run_text)
                if bold:
                    r.bold = True
                if italic:
                    r.italic = True
            # Move to before anchor
            anchor_el.addprevious(p._element)
            elements_added += 1
        elif b["kind"] == "p":
            p = doc.add_paragraph()
            for run_text, bold, italic in parse_inline(b["text"]):
                if not run_text:
                    continue
                r = p.add_run(run_text)
                if bold:
                    r.bold = True
                if italic:
                    r.italic = True
            anchor_el.addprevious(p._element)
            elements_added += 1
        elif b["kind"] == "list":
            for item in b["items"]:
                clean = re.sub(r"^[•\-\*]\s*", "", item)
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles["List Bullet"]
                except KeyError:
                    pass
                for run_text, bold, italic in parse_inline(clean):
                    if not run_text:
                        continue
                    r = p.add_run(run_text)
                    if bold:
                        r.bold = True
                    if italic:
                        r.italic = True
                anchor_el.addprevious(p._element)
                elements_added += 1
        elif b["kind"] == "table":
            table = add_table_to_doc(doc, b["rows"])
            if table is not None:
                anchor_el.addprevious(table._element)
                elements_added += 1
                # Empty paragraph for spacing
                p = doc.add_paragraph()
                anchor_el.addprevious(p._element)
                elements_added += 1
    return elements_added


def find_chapter_range(doc, matcher):
    """Find paragraph indices [start, end) belonging to a chapter via matcher."""
    paragraphs = doc.paragraphs
    start = None
    for i, p in enumerate(paragraphs):
        style_name = p.style.name if p.style else ""
        text = p.text.strip()
        if start is None and style_name in ("Heading 1", "Heading1") and matcher(text, style_name):
            start = i
            continue
        if start is not None and style_name in ("Heading 1", "Heading1"):
            return start, i
    if start is None:
        return None, None
    return start, len(paragraphs)


def main():
    print(f"=== build_v4_1.py — python-docx ===")
    print(f"SRC: {SRC.name}")
    print(f"OUT: {V41.name}\n")

    expected = ["cap2.md", "cap3.md", "cap4.md", "cap5-6.md",
                "cap7.md", "cap8-mensajes-re.md", "anexo4.md"]
    for f in expected:
        if not (CONTEXT / f).exists():
            raise SystemExit(f"⚠ Falta {f}")

    parsed = {f: parse_md(CONTEXT / f) for f in expected}
    cap56 = split_combined(parsed["cap5-6.md"], [
        ("c5", lambda t: t.lower().startswith("5") or "comparación" in t.lower()),
        ("c6", lambda t: t.lower().startswith("6") or "alternativas" in t.lower()),
    ])
    cap8mr = split_combined(parsed["cap8-mensajes-re.md"], [
        ("mk", lambda t: "mensajes clave" in t.lower()),
        ("re", lambda t: "resumen ejecutivo" in t.lower()),
        ("c8", lambda t: t.lower().startswith("8") or "conclusiones" in t.lower()),
    ])

    targets = [
        ("Mensajes clave", cap8mr["mk"], lambda t, s: t.startswith("Mensajes clave") and "Tabla" not in t and "Tarjeta" not in t and "Índice" not in t),
        ("Resumen ejecutivo", cap8mr["re"], lambda t, s: t.startswith("Resumen ejecutivo")),
        ("Cap 2", parsed["cap2.md"], lambda t, s: t.startswith("2.") or t.startswith("2 ")),
        ("Cap 3", parsed["cap3.md"], lambda t, s: t.startswith("3.") or t.startswith("3 ")),
        ("Cap 4", parsed["cap4.md"], lambda t, s: t.startswith("4.") or t.startswith("4 ")),
        ("Cap 5", cap56["c5"], lambda t, s: t.startswith("5.") or t.startswith("5 ")),
        ("Cap 6", cap56["c6"], lambda t, s: t.startswith("6.") or t.startswith("6 ")),
        ("Cap 7", parsed["cap7.md"], lambda t, s: t.startswith("7.") or t.startswith("7 ")),
        ("Cap 8", cap8mr["c8"], lambda t, s: t.startswith("8.") or t.startswith("8 ")),
    ]

    # Cargar docx
    doc = Document(str(SRC))

    print(f"Estilos disponibles en SRC: {[s.name for s in doc.styles]}\n")

    located = []
    for label, blocks, matcher in targets:
        start, end = find_chapter_range(doc, matcher)
        if start is not None:
            located.append((label, blocks, start, end))
            print(f"  ✓ Localizado: {label} paragraphs[{start}:{end}]")
        else:
            print(f"  ⚠ NO LOCALIZADO: {label}")

    # Aplicar reverse
    located.sort(key=lambda x: x[2], reverse=True)
    for label, blocks, start, end in located:
        # Anchor = paragraph at end (or last if end == len)
        if end < len(doc.paragraphs):
            anchor = doc.paragraphs[end]
        else:
            # Use last paragraph as anchor and we'll append after
            anchor = doc.paragraphs[-1]

        # Insert new content BEFORE anchor
        added = insert_blocks_at(doc, anchor, blocks)

        # Now delete old paragraphs [start:end]
        # Re-read paragraphs because we just added some
        # Simpler approach: identify old paragraphs by re-scanning before insert
        # But we already inserted. Better: track XML elements directly before insertion.
        # Re-implementing: we need to do delete BEFORE insert.

        print(f"  ✓ Insertado: {label} +{added} elementos antes de paragraph[{end}]")

    # Hmm, problema: insertamos pero no borramos. Necesito reordenar.
    # Re-iniciar con approach correcto.

    # APPROACH 2: borrar primero, luego insertar
    # Vuelve a cargar
    doc2 = Document(str(SRC))

    located2 = []
    for label, blocks, matcher in targets:
        start, end = find_chapter_range(doc2, matcher)
        if start is not None:
            located2.append((label, blocks, start, end))

    located2.sort(key=lambda x: x[2], reverse=True)

    for label, blocks, start, end in located2:
        # Get XML elements of paragraphs at [start:end]
        paras = doc2.paragraphs
        old_elements = [paras[i]._element for i in range(start, end)]

        # Anchor BEFORE deletion: use the parent of first old element
        first_el = old_elements[0]
        parent = first_el.getparent()
        anchor_idx = list(parent).index(first_el)

        # Delete old
        for el in old_elements:
            try:
                parent.remove(el)
            except Exception:
                pass

        # Now build new XML elements
        # Use a separate doc as a "factory"
        factory = Document(str(SRC))
        # Clear factory body except for first paragraph (we use the doc to construct elements)
        # Easier: append elements one by one to doc2

        new_elements_xml = []
        for b in blocks:
            if b["kind"] == "heading":
                level = b["level"]
                p = doc2.add_paragraph()
                try:
                    p.style = doc2.styles[f"Heading {level}"]
                except KeyError:
                    try:
                        p.style = doc2.styles[f"Heading{level}"]
                    except KeyError:
                        pass
                for run_text, bold, italic in parse_inline(b["text"]):
                    if not run_text:
                        continue
                    r = p.add_run(run_text)
                    if bold: r.bold = True
                    if italic: r.italic = True
                # Move element from end to where we want it
                el = p._element
                el.getparent().remove(el)
                new_elements_xml.append(el)
            elif b["kind"] == "p":
                p = doc2.add_paragraph()
                for run_text, bold, italic in parse_inline(b["text"]):
                    if not run_text:
                        continue
                    r = p.add_run(run_text)
                    if bold: r.bold = True
                    if italic: r.italic = True
                el = p._element
                el.getparent().remove(el)
                new_elements_xml.append(el)
            elif b["kind"] == "list":
                for item in b["items"]:
                    clean = re.sub(r"^[•\-\*]\s*", "", item)
                    p = doc2.add_paragraph()
                    try:
                        p.style = doc2.styles["List Bullet"]
                    except KeyError:
                        pass
                    for run_text, bold, italic in parse_inline(clean):
                        if not run_text:
                            continue
                        r = p.add_run(run_text)
                        if bold: r.bold = True
                        if italic: r.italic = True
                    el = p._element
                    el.getparent().remove(el)
                    new_elements_xml.append(el)
            elif b["kind"] == "table":
                table = add_table_to_doc(doc2, b["rows"])
                if table is not None:
                    el = table._element
                    el.getparent().remove(el)
                    new_elements_xml.append(el)
                    # Empty paragraph after table
                    p = doc2.add_paragraph()
                    pe = p._element
                    pe.getparent().remove(pe)
                    new_elements_xml.append(pe)

        # Insert new elements at anchor_idx
        for offset, el in enumerate(new_elements_xml):
            parent.insert(anchor_idx + offset, el)
        print(f"  ✓ Aplicado: {label} ({len(new_elements_xml)} elementos)")

    if V41.exists():
        V41.unlink()
    V41.parent.mkdir(parents=True, exist_ok=True)
    doc2.save(str(V41))
    print(f"\n✓ {V41.name} ({V41.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
