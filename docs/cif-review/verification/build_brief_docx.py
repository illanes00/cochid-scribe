"""Genera brief-medicamentos-v4.docx con estilos del informe v3.18.

Toma como template el v3.18 (para heredar Heading1/Heading2/Heading3/Normal),
borra todo el contenido del body, y reescribe con el contenido del brief v4.

NO usa pandoc/markdown processing. Construye párrafos con python-docx
+ XML directo cuando hace falta (para tablas, callouts).

Salida: /srv/projects/cochid/cochid-scribe/docs/cif-review/output/brief/brief-medicamentos-v4.docx
"""

from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

BASE = Path("/srv/projects/cochid/cochid-scribe/docs/cif-review")
TEMPLATE = BASE / "output/informe-final-v3.18-aceptada.docx"
OUT = BASE / "output/brief/brief-medicamentos-v4.docx"


def main():
    print("=== build_brief_docx.py ===")
    print(f"Template: {TEMPLATE.name}")
    print(f"Output: {OUT.name}\n")

    # Cargar template
    doc = Document(str(TEMPLATE))

    # Borrar todo el contenido del body (mantener sectPr al final)
    body = doc.element.body
    sectPr = None
    # Encontrar último sectPr (define márgenes etc)
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sectPr = child
            continue
        body.remove(child)
    print("Body limpiado (sectPr preservado)\n")

    # ====================================================
    # PORTADA
    # ====================================================
    # Título grande
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Inclusión sostenible de medicamentos en los planes de salud en Chile")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Síntesis para política pública")
    r.font.size = Pt(16)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    p = doc.add_paragraph()  # espacio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MINUTA TÉCNICA")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x60, 0x40, 0x80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Martín Illanes · Espacio Público")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mayo de 2026")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Investigación financiada por la Cámara de la Innovación Farmacéutica de Chile")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Salto de página
    p = doc.add_paragraph()
    p.add_run().add_break()

    # ====================================================
    # PÁGINA 1 — El panorama + Mensajes clave
    # ====================================================
    h = doc.add_paragraph(style="Heading1")
    h.add_run("El panorama")

    doc.add_paragraph(
        "El sistema chileno de salud cubre, en el papel, a la totalidad de la población a través "
        "de Fonasa y de las Isapres. La cobertura efectiva de medicamentos, sin embargo, está lejos "
        "de seguir esa universalidad. El 62% del gasto total en medicamentos lo financian directamente "
        "los hogares (CIF y Escuela de Gobierno UC, 2025); en el segmento ambulatorio retail la "
        "proporción se eleva al 71% (OECD SHA, 2022). Es la cifra más alta de la OCDE en ambos cortes "
        "y se ha mantenido estable durante la última década."
    )

    doc.add_paragraph(
        "Esta minuta sintetiza el informe extenso del mismo título. Ordena la magnitud y la naturaleza "
        "del gasto de bolsillo en medicamentos; reconstruye la tabla de protección farmacéutica vigente "
        "con sus instrumentos y sus brechas; reseña las trayectorias de cobertura de sistemas OCDE "
        "consolidados; y plantea tres escenarios de política para Chile, con sus reformas legales "
        "asociadas. No formula una recomendación única; ordena la conversación."
    )

    h = doc.add_paragraph(style="Heading2")
    h.add_run("Mensajes clave")

    mensajes = [
        ("1.", "Gasto de bolsillo estructural. 62% del gasto total en medicamentos (CIF/UC, 2024) y 71% del gasto retail farmacéutico (OECD SHA, 2022) lo financian los hogares. Es la cifra más alta de OCDE en ambos cortes."),
        ("2.", "Dos miradas complementarias. Una mirada acumulativa, donde el costo crónico se suma mes a mes (medicamentos ambulatorios de uso permanente). Una mirada catastrófica, donde un evento aislado puede agotar el patrimonio familiar (tratamientos de alto costo). Ninguna medida única atiende ambas a la vez."),
        ("3.", "Patrón de paquete OCDE. La evidencia comparada muestra que la protección efectiva no se logra con un instrumento único, sino combinando seis componentes: canasta priorizada, copagos topados, dispensación accesible, sustitución por valor, regulación focal y trazabilidad. Los países que reducen sostenidamente el OOP integran al menos cuatro de los seis."),
        ("4.", "BFU como zoom analítico. El informe profundiza uno de los tres escenarios analizados, el de convergencia intermedia con el cluster OCDE. Su eje estructural es un Beneficio Farmacéutico Universal (BFU) aplicable a beneficiarios de Fonasa e Isapre bajo reglas comunes. La elección no es prescriptiva: BFU concentra decisiones de diseño que iluminan también E1 (ajuste gradual) y E3 (convergencia plena)."),
        ("5.", "Transparencia presupuestaria pendiente. El presupuesto público de medicamentos GES en Fonasa, GES en Isapre y la porción farmacéutica de la CAEC no se publican de forma desagregada. La rendición de cuentas y el monitoreo del gasto agregado requieren cerrar esa brecha de transparencia."),
    ]
    for num, text in mensajes:
        p = doc.add_paragraph()
        r = p.add_run(num + " ")
        r.font.bold = True
        p.add_run(text)

    # ====================================================
    # PÁGINA 2 — Lógica acumulativa
    # ====================================================
    p = doc.add_paragraph()
    p.add_run().add_break()

    h = doc.add_paragraph(style="Heading1")
    h.add_run("El gasto cotidiano: la lógica acumulativa")

    doc.add_paragraph(
        "Los medicamentos para enfermedades crónicas de alta prevalencia, como hipertensión arterial, "
        "diabetes mellitus tipo 2, dislipidemia, asma y depresión, concentran una proporción sustantiva "
        "del gasto farmacéutico de los hogares chilenos. La razón estandarizada de comparación "
        "internacional es la proporción del gasto retail farmacéutico financiada por los hogares "
        "(OECD SHA, ratio HF3 sobre HC51). Para Chile esta razón se ubicó en 71% en 2022 y se ha "
        "mantenido entre 70% y 72% durante el período 2019 a 2023. El promedio OCDE para la misma "
        "razón se sitúa en torno al 39%."
    )

    doc.add_paragraph(
        "La encuesta nacional Ipsos y Espacio Público de julio de 2025 documenta el correlato individual "
        "de esa cifra agregada: 29% de los consultados reportó haber dejado de tomar alguna dosis de "
        "medicamento en los doce meses previos por su costo, y entre quienes lo hicieron, 70% lo hizo "
        "en más de una oportunidad. La interrupción de tratamiento por motivos económicos no es un "
        "riesgo teórico, sino una práctica documentada que afecta al control de patologías crónicas "
        "y que se manifiesta con mayor intensidad en hogares de ingresos medios y bajos, en personas "
        "mayores y en residentes de comunas con menor cobertura de farmacia comunitaria pública."
    )

    h = doc.add_paragraph(style="Heading3")
    h.add_run("Algunas magnitudes relevantes")
    bullets = [
        "Gasto de bolsillo sobre gasto total en medicamentos, Chile, 2024: 62% (CIF y UC, 2025).",
        "Gasto de bolsillo sobre gasto retail farmacéutico, Chile, 2022: 71% (OECD SHA, HF3 sobre HC51).",
        "Gasto público en medicamentos como porcentaje del PIB: 0,46% total (CIF y UC, 2024); 0,37% retail (OECD SHA, 2023). Promedio OCDE total cercano al 1,5% del PIB.",
        "Tope anual de gasto de bolsillo en medicamentos ambulatorios: inexistente tanto en Fonasa como en Isapres.",
        "Hogares que reportan haber suspendido alguna dosis por costo en el último año: 29% (Ipsos y Espacio Público, julio 2025).",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="ListBullet")
        p.add_run(b)

    doc.add_paragraph(
        "La caracterización del problema como acumulativo, antes que como problema de precio unitario, "
        "tiene una consecuencia operativa relevante. Un copago razonable para una compra esporádica deja "
        "de serlo cuando se replica mes a mes durante años. La ausencia de un tope anual en el sistema "
        "chileno, instrumento estándar en los países OCDE de referencia, es la brecha más visible que "
        "separa a Chile del cluster comparativo."
    )

    # ====================================================
    # PÁGINA 3 — Alto costo + tabla protección
    # ====================================================
    p = doc.add_paragraph()
    p.add_run().add_break()

    h = doc.add_paragraph(style="Heading1")
    h.add_run("Alto costo y cobertura fragmentada")

    doc.add_paragraph(
        "Las terapias de alto costo plantean una segunda lógica de riesgo. En Chile el principal "
        "instrumento es la Ley Ricarte Soto (Ley 20.850 de 2015), que cubre 27 patologías de alto "
        "costo definidas por decreto y otorga cobertura financiera integral a quienes cumplen los "
        "criterios clínicos, con un acumulado de 65.351 personas atendidas a 2024. El presupuesto "
        "regular de la Ley Ricarte Soto alcanzó MM$175.672 en 2025 (Dipres, Ley de Presupuestos). "
        "Por diseño, su perímetro es acotado y opera como mecanismo de excepción para riesgo "
        "catastrófico, no como beneficio ambulatorio general."
    )

    doc.add_paragraph(
        "La tabla de protección farmacéutica chilena combina seis instrumentos con criterios de "
        "elegibilidad heterogéneos. Sus perímetros se solapan parcialmente y dejan brechas que "
        "explican el volumen del gasto de bolsillo."
    )

    # Tabla protección
    table = doc.add_table(rows=7, cols=3)
    table.style = "TableGrid"
    headers = ["Instrumento", "Beneficiarios", "Cobertura"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True

    rows = [
        ("GES (medicamentos)", "Fonasa + Isapre", "87 problemas; sin glosa segregada en presupuesto"),
        ("Ley Ricarte Soto", "Fonasa + Isapre", "27 patologías de alto costo (Decreto 4 de 2019)"),
        ("DAC (Glosa 11)", "Fonasa", "Oncológicos no cubiertos por GES ni LRS, caso a caso"),
        ("FOFAR + Arsenal APS", "Fonasa en APS", "Esenciales para hipertensión, diabetes, dislipidemia"),
        ("CAEC (seguro adicional Isapre)", "Isapre", "Catastrófico privado; porción farmacéutica no publicada"),
        ("Cenabast Ley 21.198", "Hogares vía retail adherido", "Subsidio cruzado vía precio en farmacias adheridas"),
    ]
    for i, (inst, ben, cob) in enumerate(rows, 1):
        table.cell(i, 0).text = inst
        table.cell(i, 1).text = ben
        table.cell(i, 2).text = cob

    p = doc.add_paragraph()
    p.add_run("Esa fragmentación, antes que la ausencia de instrumentos, explica buena parte del gasto "
             "de bolsillo persistente. Los recursos existen, pero no convergen en un beneficio coherente "
             "para el paciente. La judicialización vía recursos de protección representó cerca de "
             "MM$81.000 en 2024 (Fonasa, cuenta pública), equivalente al 5,3% del gasto público en "
             "medicamentos. Su crecimiento sostenido sugiere que la vía judicial está supliendo, en "
             "una fracción significativa de casos, la ausencia de un canal ordinario de incorporación.")

    # ====================================================
    # PÁGINA 4 — Comparado
    # ====================================================
    p = doc.add_paragraph()
    p.add_run().add_break()

    h = doc.add_paragraph(style="Heading1")
    h.add_run("Lo que muestra la evidencia comparada")

    doc.add_paragraph(
        "El informe analiza 13 países OCDE con cobertura farmacéutica universal o cuasi-universal, "
        "datos disponibles en OECD Health Statistics y diversidad institucional suficiente para "
        "iluminar el debate chileno. La comparación se ordena en tres clusters fiscales según el "
        "esfuerzo público en medicamentos como porcentaje del PIB."
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = "TableGrid"
    th_rows = [
        ("Cluster", "Países", "Gasto público (% PIB)"),
        ("Alta cobertura", "Alemania, Francia, Países Bajos", "1,3% a 1,5%; OOP por debajo del 25%"),
        ("Cluster intermedio OCDE", "España, Reino Unido, Canadá", "0,8% a 1,2%; OOP entre 26% y 43%"),
        ("Bajo esfuerzo fiscal", "Chile", "0,46% total; OOP entre 62% y 71%"),
    ]
    for i, (a, b, c) in enumerate(th_rows):
        table.cell(i, 0).text = ""
        table.cell(i, 1).text = ""
        table.cell(i, 2).text = ""
        for j, val in enumerate([a, b, c]):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if i == 0:
                r.font.bold = True

    doc.add_paragraph(
        "Los sistemas con bajo gasto de bolsillo en medicamentos no comparten una arquitectura única. "
        "Comparten la operación simultánea de varios elementos combinables: una canasta explícita "
        "evaluada por una agencia técnica, copagos calibrados con tope financiero anual, un modelo "
        "de dispensación con convenio que asegura continuidad terapéutica, sustitución por "
        "bioequivalente con incentivo de precio, regulación focalizada cuando no hay competencia "
        "y trazabilidad de prescripción y consumo. Donde uno de esos elementos falla, el sistema "
        "no logra sostener la cobertura efectiva."
    )

    # ====================================================
    # PÁGINA 5 — Tres escenarios
    # ====================================================
    p = doc.add_paragraph()
    p.add_run().add_break()

    h = doc.add_paragraph(style="Heading1")
    h.add_run("Tres escenarios de política para Chile")

    doc.add_paragraph(
        "El informe construye tres escenarios alternativos con sus reformas legales y trade-offs "
        "fiscales asociados. Las cifras son órdenes de magnitud calibrados sobre el PIB chileno 2024 "
        "(en precios 2022; orden de magnitud)."
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = "TableGrid"
    rows = [
        ("Escenario", "Costo fiscal anual", "OOP retail esperado"),
        ("E1 Ajuste gradual", "USD 400 a 500 millones", "Reducción a 60%"),
        ("E2 BFU intermedio", "USD 800 a 900 millones", "Reducción al rango 30 a 40%"),
        ("E3 Convergencia plena OCDE", "USD 2.500 a 3.000 millones", "Niveles bajos OCDE"),
    ]
    for i, (a, b, c) in enumerate(rows):
        for j, val in enumerate([a, b, c]):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if i == 0:
                r.font.bold = True

    doc.add_paragraph(
        "El primer escenario combina expansión moderada de la Ley Ricarte Soto a más patologías y "
        "un techo anual de bolsillo sobre los instrumentos vigentes. El segundo introduce un "
        "Beneficio Farmacéutico Universal articulado con los regímenes existentes. El tercero "
        "lleva el sistema chileno al cuartil superior OCDE, con copagos casi simbólicos y canasta "
        "expandida; supone un esfuerzo fiscal y una capacidad institucional que excede el horizonte "
        "presupuestario inmediato."
    )

    # ====================================================
    # PÁGINA 6 — Zoom BFU
    # ====================================================
    p = doc.add_paragraph()
    p.add_run().add_break()

    h = doc.add_paragraph(style="Heading1")
    h.add_run("Zoom analítico: Beneficio Farmacéutico Universal")

    doc.add_paragraph(
        "El Capítulo 7 del informe desarrolla en detalle el Escenario 2 a través del Beneficio "
        "Farmacéutico Universal (BFU). El BFU opera como red de seguridad financiera para el "
        "gasto de bolsillo residual del hogar en medicamentos, articulado con los regímenes "
        "específicos vigentes (GES, Ley Ricarte Soto, DAC, FOFAR, CAEC) y aplicable a beneficiarios "
        "de Fonasa e Isapre bajo reglas comunes."
    )

    h = doc.add_paragraph(style="Heading3")
    h.add_run("Componentes del BFU")
    components = [
        "Lista positiva universal priorizada por evaluación de tecnologías sanitarias (ETESA), con criterios de evidencia clínica, costo-efectividad y carga de enfermedad.",
        "Tope anual de gasto de bolsillo del hogar, calibrado en 13% a 15% del ingreso per cápita (USD 800 a 900 millones anuales según simulación EPF, anexo metodológico §5).",
        "Modalidad de subsidio en punto de venta para reducir barreras de acceso en quintiles 1 a 3.",
        "Sustitución bioequivalente obligatoria con incentivo de precio.",
        "Cobertura sobre todo el gasto del hogar en medicamentos, retail y hospitalario ambulatorio, sin distinción de canal.",
        "Trazabilidad por RUT con continuidad entre Fonasa e Isapre.",
        "Acuerdos de acceso gestionado y mecanismos de riesgo compartido para innovación de alto costo.",
    ]
    for c in components:
        p = doc.add_paragraph(style="ListBullet")
        p.add_run(c)

    doc.add_paragraph(
        "El BFU no reemplaza GES, Ley Ricarte Soto ni DAC. Actúa como red de protección sobre el "
        "gasto residual cuando esos regímenes no cubren, evitando duplicación y manteniendo la "
        "precedencia de los instrumentos garantizados por ley para sus respectivas poblaciones objetivo."
    )

    # ====================================================
    # PÁGINA 7 — 5 preguntas técnicas para la discusión
    # ====================================================
    p = doc.add_paragraph()
    p.add_run().add_break()

    h = doc.add_paragraph(style="Heading1")
    h.add_run("Cinco preguntas técnicas para la discusión")

    doc.add_paragraph(
        "El informe deja cinco preguntas abiertas que el debate público y el seminario buscan "
        "ordenar. Su resolución corresponde a la deliberación entre actores con preferencias y "
        "diagnósticos distintos."
    )

    preguntas = [
        ("Fuente unificada para fichas país.", "Los Anexos 6 y 7 muestran datos por país en metodologías distintas a la Tabla 4 del Capítulo 5. Convendría unificar sobre OECD SHA, dejando CIF/UC como referencia complementaria de agregado fiscal."),
        ("Articulación operativa BFU + GES + LRS + DAC + FOFAR.", "El Capítulo 7.3 define la precedencia conceptual. La discusión pendiente es la regla de aplicación caso a caso: si GES cubre la patología, ¿el copago GES vigente prevalece y el BFU cubre solo la diferencia residual? ¿Qué sucede cuando un fármaco está en GES para una indicación y se prescribe para otra no garantizada?"),
        ("Unidad del tope: persona u hogar.", "El Capítulo 7.4.2 plantea ambas alternativas. Persona es simple administrativamente y reconoce trayectoria clínica individual; hogar captura mejor el riesgo financiero agregado pero introduce dificultades de definición. La elección define la regresividad efectiva del beneficio."),
        ("Clasificación de medidas por horizonte de reforma.", "El Capítulo 8.1 propone una tabla de corto, mediano y largo plazo. Las medidas que requieren acuerdo legislativo (creación de la canasta universal, gobernanza de la agencia evaluadora) se distinguen de las que pueden implementarse vía decreto o reglamento."),
        ("CENABAST como compras agregadas dentro del beneficio.", "El Capítulo 7.6 plantea CENABAST como agregador de demanda con la Ley 21.198 ya vigente. La pregunta es si el alcance de la Ley 21.198 puede expandirse al beneficio farmacéutico universal o si se requiere una nueva habilitación legal."),
    ]
    for titulo, texto in preguntas:
        p = doc.add_paragraph()
        r = p.add_run(titulo + " ")
        r.font.bold = True
        p.add_run(texto)

    # ====================================================
    # PÁGINA 8 — Nota metodológica
    # ====================================================
    p = doc.add_paragraph()
    p.add_run().add_break()

    h = doc.add_paragraph(style="Heading1")
    h.add_run("Nota metodológica")

    doc.add_paragraph(
        "Este documento sintetiza el informe extenso producido por Espacio Público con financiamiento "
        "de la Cámara de la Innovación Farmacéutica de Chile. Las posiciones expresadas son "
        "responsabilidad de Espacio Público. El proceso editorial incluyó revisión metodológica de "
        "Carla Castillo (UDD) y Eduardo Undurraga (PUC)."
    )

    doc.add_paragraph(
        "Las cifras de gasto público en medicamentos se construyen a partir de tres fuentes "
        "complementarias: el sistema OECD System of Health Accounts (HC51 retail), el estudio "
        "Caracterización del Gasto Público en Medicamentos (CIF y Escuela de Gobierno UC, segunda "
        "edición, octubre 2025) y la Encuesta de Presupuestos Familiares IX (INE, 2021-2022). La "
        "convergencia entre las tres fuentes en el orden 0,46% del PIB confirma la robustez del "
        "diagnóstico cuantitativo. La metodología detallada se desarrolla en el Anexo 4 del informe extenso."
    )

    p = doc.add_paragraph()
    r = p.add_run("Versión v4 · 7 de mayo de 2026")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Re-add sectPr at end
    if sectPr is not None:
        body.append(sectPr)

    # Save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"\n✓ {OUT.name} ({OUT.stat().st_size:,} bytes)")
    print(f"  Páginas estimadas: 8")
    print(f"  Estilos heredados de v3.18: Heading1, Heading2, Heading3, Normal, ListBullet")


if __name__ == "__main__":
    main()
